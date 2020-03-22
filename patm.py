# # alternetive solution with just typing answer next to quesrion
import docx
import random
from docx.shared import Pt

from docx.shared import RGBColor

patm = docx.Document(
    r"C:\Users\Hayk\Desktop\patm qnnutun patasxannerov\3-րդ շտեմարան\Հարցեր.docx")

quant = len(patm.paragraphs)

all_questions = []

def get_indexes_of_questions():
    """
    gets indexes of all the questions
    """
    print ("starting getting indexes")
    for i in range(quant):
        if i % 100 == 0:
            print(i)
        if patm.paragraphs[i].style.name == "Heading 4":
            all_questions.append(i)
    return all_questions


# all_questions = get_indexes_of_questions()
# all_answers = ['3', '3', '4', '2', '2', '3', '2', '2', '4', '1', '2', '4', '1', '4', '3', '2', '2', '3', '2', '2', '4', '2', '2', '1', '3', '2', '2', '3', '3', '2', '1', '3', '4', '3', '1', '2', '1', '3', '3', '3', '2', '2', '1', '2', '4', '1', '2', '2', '1', '4', '2', '2', '2', '4', '4', '4', '3', '2', '3', '4', '2', '2', '4', '3', '3', '4', '2', '3', '3', '2', '3', '1', '2', '4', '3', '4', '4', '1', '3', '3', '2', '3', '2', '4', '2', '2', '3', '2', '2', '2', '3', '2', '3', '3', '3', '1', '3', '1', '4', '2', '3', '2', '2', '2', '3', '2', '4', '2', '1', '4', '3', '4', '2', '3', '1', '4', '2', '3', '3', '2', '1', '4', '2', '4', '2', '2', '4', '3', '2', '2', '1', '1', '1', '1,3', '3', '3, 6, 7', '7, 8, 9', '2, 6, 8', '4, 5, 6', '1, 5, 7', '2, 4, 5', '1, 4, 7', '2, 5, 8', '3, 6, 7', '2, 3, 8', '1, 2, 7', '3, 4', '4', '1', '3', '4', '3', '1', '4', '4', '3', '3', '2', '4', '1', '3', '1, 4', '2, 3', '3, 5', '1', '1', '3', '4', '3', '1', '4', '2', '3', '4', '5, 7', '3', '3', '4', '4', '3', '4', '2', '3', '3', '4', '4,6', '5', '3, 6', '2, 4', '2, 6', '2', '2', '3', '1', '4', '2', '4', '1', '4', '3', '1', '4', '2', '4', '2', '3', '3', '2', '2', '3', '4', '1', '4', '2', '1', '3', '1', '3', '2', '3, 6, 7', '2', '1', '3', '3', '3', '2', '4', '3', '2', '2', '2', '3, 4', '2', '1', '3', '1, 3, 8', '3, 4, 6', '1, 5', '1,5', '1, 4', '1, 4, 5', '1, 5, 8', '3, 5', '1, 5, 7', '1', '2, 5, 7', '3, 4', '2', '1, 3', '4', '3', '3', '3', '4', '1', '4', '2', '1', '2', '4', '2', '2', '2', '2', '2', '3', '2', '4', '2', '3', '4', '2', '1', '3', '1', '4', '2', '4', '2', '3', '1', '3', '2', '3', '4', '4', '4', '2', '3', '1', '3', '4', '4', '2', '3', '2', '3', '3', '2', '3', '3', '3', '2', '2', '1', '2', '3', '3', '3', '1', '4', '1', '2', '2', '1', '1', '2', '4', '2', '4', '4', '2', '4', '1', '3', '3', '3', '2', '2', '4', '4', '3', '3', '4', '3', '3', '3', '2', '2', '4', '3', '1', '2', '3', '3', '3', '3', '4', '2', '3', '3', '2', '3', '2', '2', '1', '1', '4', '1', '1', '4', '4', '3', '4', '2', '1', '2', '1', '4', '3', '4', '4', '4', '2', '1', '1', '3', '1', '2', '4', '2', '2', '3', '1', '1', '4', '1', '2', '3', '4', '2', '4', '1', '2', '1', '3', '4', '1', '3', '2', '3', '1', '2', '2', '2', '3', '3', '2', '3', '1', '3', '2', '3', '2', '2', '4', '4', '3', '2', '2', '4', '4', '4', '1', '2', '1', '4', '1', '3', '4', '1', '1', '2', '1', '2', '1', '2', '3', '1', '2', '1', '4', '2', '4', '2', '2', '1', '2', '3', '2', '2', '4', '4', '3', '4', '4', '2', '2', '3', '3', '3', '2', '3', '4', '3', '4', '3', '2', '2', '2', '1', '2', '3', '3', '1', '4', '4', '3', '1', '3', '4', '4', '2', '1', '4', '4', '4', '2', '1', '1', '2', '1', '2', '2', '1', '1', '3, 5', '3, 7, 8', '4', '3', '2', '2', '1', '3', '1', '1', '3', '1', '1', '4', '1', '2', '3', '2, 5', '1, 4', '2, 3', '1, 4', '1, 5, 6', '4', '4', '2', '2', '4', '4', '2', '3', '4', '2', '1, 5', '1', '3', '2', '1', '4', '2', '2', '3', '1', '1, 4, 6', '3', '2', '2', '1', '2, 3, 5', '3', '1', '2', '1', '4', '2', '3', '4', '3', '2', '2', '2', '3', '3', '4', '1', '4', '4', '4', '2', '3', '4', '2', '1', '4', '1', '2', '4', '1', '3', '4', '3', '4,6', '4', '5, 8', '1', '5, 9', '4', '3,7', '3', '2', '4, 6', '3', '2', '2', '3', '3', '1', '3', '2', '1', '2', '2', '3', '1', '3', '3', '1', '4', '2', '1', '2', '3', '2', '3', '2', '4', '3', '3', '2', '4', '3', '3', '1', '4', '1', '1', '2', '2', '4', '1', '4', '4', '1', '3', '2', '4', '3', '3', '1', '3', '3', '2', '4', '1', '1', '1', '2', '1', '2', '1', '3', '2', '4', '2', '3', '1', '3', '3', '4', '3', '2', '3', '4', '4', '2', '1', '2', '3', '2', '4', '1', '4', '1', '4', '1,5', '2, 4, 7', '1,5', '3', '1', '1', '3', '2', '3', '3', '4', '3', '2', '2', '4', '3', '1, 3', '2', '4', '1', '3, 6', '3, 5', '2', '3', '2, 3, 6', '3', '3,5', '3', '4', '3', '3, 5, 7', '1', '1', '2, 4', '4, 5, 7', '3, 4, 7', '2,5', '2', '2', '2, 4', '3', '1, 5', '1', '2', '2', '1', '1', '4', '1', '1', '4', '3', '2, 3, 5', '1, 4, 6', '2, 5', '2', '1, 3', '2, 3, 6', '3, 4, 6', '4', '2, 4', '3', '4', '3', '1', '2', '1', '4', '4', '2', '3', '2', '1', '1', '1', '2', '1', '2', '1', '3', '1', '3', '1', '1', '3', '1', '2', '1', '3', '3', '4', '1', '3', '2', '3', '4', '3', '4', '1', '1', '2', '2', '1', '2', '1', '1', '3', '3', '3', '3', '1', '1', '1', '1', '1', '1', '3', '1', '1', '4', '2', '3', '4', '1', '2', '2', '3', '3', '1', '2', '2', '4', '4', '3', '2', '2', '3', '2', '3', '1', '3', '2', '3', '1', '2', '1', '1', '1', '1', '2', '3', '1', '2', '2', '1', '4', '1', '1', '3', '1', '1', '2', '1', '2', '4', '1', '2', '1', '2', '2', '1', '3', '3', '2', '4', '3', '2', '2', '3', '4', '3', '3', '1', '1', '1', '1', '1', '2', '3', '3', '1', '1', '4', '1', '1', '4', '3', '3', '3', '2', '2', '3', '2', '3', '1', '1', '3', '3', '1', '1', '2', '3', '4', '2, 5', '1', '2', '3', '3', '2', '3', '3', '1', '4', '1', '1', '2', '3', '2', '2', '3', '3', '3', '2', '3', '2', '2', '2', '3', '2', '2', '3', '2, 4', '1', '1', '1, 3', '3', '3', '4', '1', '3', '2', '1', '1', '1', '3', '2', '2', '1', '2', '4', '3', '3', '4', '3', '2', '2', '3', '2', '4', '4', '2', '2', '1', '4', '3', '3, 5', '1', '4', '2', '2,4', '2', '3', '1', '2', '4', '1', '2', '4', '1', '4', '3', '3', '2', '2', '3', '3', '5,7', '4', '4,6', '4,6', '1,5', '2,6', '3', '2', '3', '2,6', '4', '1', '4', '3', '2', '1', '4', '3', '4', '2', '1', '1', '1', '1', '2', '1', '1', '5', '2', '1', '1', '1', '1', '3', '2', '2', '2', '3', '4', '3', '1', '1', '2', '2', '3', '4', '4', '1', '1', '4', '1', '5', '4', '4', '2', '1', '2', '2', '3', '3,7', '3,7', '3,9', '2', '2', '3', '3', '4', '2', '2', '1', '1', '3', '4', '2', '1', '1', '1', '1', '1', '3', '1', '1', '4', '2', '3', '3', '1', '3', '3', '1', '1', '3', '3', '3', '3', '3', '4', '3', '3', '3', '4', '1', '2', '3', '3', '3', '2', '1', '1', '2', '2', '3', '4', '3', '3', '4', '1', '3', '3', '4', '4', '1', '3', '3', '1', '1, 5, 7', '1', '3', '1', '2', '1, 3, 7', '3, 5, 7', '3', '2', '2, 5, 8', '2, 4, 6', '3, 6, 8', '1, 3, 6', '2, 6', '4', '4, 6', '3', '1,3', '2,4', '1, 4, 7', '1', '4', '3', '3', '1', '3', '4', '2', '2', '2', '1', '1', '2, 4, 7', '3, 5, 9', '1, 3, 4', '1, 3, 6', '1', '1', '1, 3, 4', '1, 4, 5', '2, 3', '1', '2', '1, 5', '2', '3', '3', '1', '4, 5', '1, 2', '1, 2, 6', '2', '3', '2, 3, 6', '2', '2', '2, 3', '3, 4', '3, 5', '4', '1, 2, 3', '1, 2, 5', '1, 3, 6', '1,2', '1', '3', '1', '2', '2', '1', '4', '2', '1', '1', '3', '3', '1', '2', '4', '2', '1', '1', '3', '3', '4', '2', '3', '1', '4', '2', '4', '4', '2', '4', '2', '4', '2', '2', '2', '2', '2', '2', '4', '2', '4', '2', '2', '2', '3', '2', '4', '3', '2', '3', '2', '1', '2', '2', '3', '3', '1', '2', '2', '1', '2', '2', '2', '3', '3', '3', '3', '1', '2', '3', '2', '3', '2', '4', '3', '3', '2', '1', '3', '3', '4', '2', '2', '2', '3', '2', '2', '1', '4', '3', '4', '3', '3', '3', '3', '2', '2', '2', '3', '4', '2', '2', '2', '2', '3', '2', '4', '2', '3', '3', '2', '3', '2', '3', '2', '1', '2', '2', '4', '4', '2', '3', '3', '3', '2', '3', '1', '3', '3', '1', '2', '4', '1', '1', '1', '3', '2', '4', '2', '3', '3', '3', '2', '1', '3', '3', '3', '3', '3', '2', '3', '4', '3', '3', '3', '2', '3', '2', '2', '2', '2', '2', '3', '2', '3', '3', '1', '3', '2', '3', '3', '2', '1', '2', '1', '4', '3', '2', '3', '3', '2', '2', '3', '4', '3', '3', '2', '2', '2', '2', '2', '4', '2', '2', '2', '2', '4', '4', '3', '2', '3', '3', '4', '2', '4', '2', '1', '2', '2', '3', '2', '1', '2', '2', '3', '2', '3', '3', '3', '2', '2', '2', '2', '2', '4', '2', '3', '2', '2', '3', '3', '1', '2', '3', '4', '1', '3', '2', '2', '3', '3', '2', '4', '3', '3', '2', '1', '4', '1', '1', '3', '1', '1', '4', '1', '2', '3', '2', '3', '3', '2', '2', '1, 5', '1, 6', '3', '3', '3', '2', '3', '1', '1', '3', '1', '3', '1, 4', '3', '4', '1', '1', '2', '4', '2', '4', '1', '2', '2', '3', '2', '4', '2', '3', '3', '2', '3', '4', '2', '3', '3', '1', '1', '4', '2', '2', '3,5,7', '3', '4', '4', '2,4', '4', '2', '3', '3', '1,3', '2', '3', '2,4', '1', '1', '1', '1', '4,5', '3,5', '3', '3', '1,4,7', '1', '4', '3', '2', '4', '3', '3', '2', '3', '2', '3', '3', '2', '3', '3', '2', '3', '4', '3', '1', '2', '2', '1', '3', '1', '3', '2', '3', '1', '3', '1', '3', '4', '1', '2', '3,7', '6,7', '4', '1', '4', '1', '2', '2', '3', '3', '2', '2,5', '4', '4', '3', '3', '2', '1', '3', '4', '4', '4', '2', '4', '4', '4', '2', '2', '3', '1', '3', '3', '1', '2', '1', '1', '2', '1', '1', '1', '1', '2', '4', '1, 3, 4', '2', '2', '2, 5, 7', '2', '1', '2', '3, 5, 7', '3', '2', '1', '4', '3', '2', '1', '2', '3', '1, 5, 7', '2', '4', '2', '2', '2', '3', '2', '2, 3, 5', '2, 4, 9', '2, 4, 7', '2, 3, 5', '1', '1', '2, 6, 8', '1, 3, 5', '1, 4, 6', '1', '3, 5', '1', '2, 4, 6', '1, 5, 6', '3, 4', '2, 3', '2, 5', '1, 3', '2', '1, 4, 7', '3', '1, 2', '2', '2, 4, 7', '2, 4', '1', '3', '4', '3', '2', '4', '3', '4', '6, 7, 9', '3', '1', '3, 4', '4', '2, 3, 6', '1', '2, 3', '1', '1, 3, 4']
# all_questions = [5, 8, 11, 14, 17, 20, 22, 24, 27, 31, 34, 37, 40, 43, 45, 48, 51, 54, 56, 59, 63, 66, 69, 72, 75, 78, 81, 84, 91, 93, 95, 97, 99, 101, 103, 108, 111, 113, 116, 118, 121, 124, 129, 132, 135, 138, 141, 143, 146, 148, 151, 154, 156, 158, 163, 168, 170, 173, 175, 178, 180, 184, 187, 189, 192, 194, 197, 200, 203, 206, 208, 210, 214, 216, 218, 220, 223, 225, 227, 230, 235, 237, 240, 246, 249, 251, 254, 257, 260, 262, 264, 267, 269, 272, 274, 278, 281, 283, 285, 287, 289, 291, 294, 297, 300, 303, 307, 310, 313, 318, 327, 331, 335, 339, 344, 349, 353, 357, 361, 365, 370, 374, 378, 386, 389, 392, 395, 398, 401, 404, 409, 415, 420, 423, 433, 438, 446, 457, 467, 476, 485, 494, 503, 513, 523, 532, 541, 547, 552, 557, 563, 568, 573, 578, 583, 588, 594, 599, 604, 609, 614, 620, 626, 632, 639, 644, 654, 660, 667, 672, 679, 688, 694, 700, 708, 718, 724, 735, 742, 752, 758, 767, 774, 781, 792, 796, 806, 816, 827, 837, 844, 852, 860, 867, 883, 892, 899, 903, 910, 917, 926, 935, 942, 951, 963, 971, 981, 991, 1001, 1006, 1016, 1022, 1028, 1039, 1056, 1069, 1079, 1082, 1087, 1095, 1100, 1106, 1111, 1116, 1121, 1126, 1131, 1137, 1142, 1147, 1152, 1158, 1164, 1169, 1178, 1188, 1196, 1203, 1209, 1215, 1224, 1234, 1240, 1249, 1252, 1260, 1266, 1272, 1278, 1296, 1298, 1301, 1304, 1307, 1309, 1311, 1313, 1315, 1317, 1319, 1322, 1325, 1327, 1329, 1331, 1334, 1336, 1338, 1340, 1342, 1344, 1346, 1349, 1351, 1354, 1356, 1358, 1360, 1363, 1365, 1367, 1369, 1371, 1373, 1375, 1378, 1380, 1383, 1385, 1388, 1390, 1392, 1394, 1396, 1398, 1400, 1402, 1404, 1406, 1409, 1411, 1414, 1416, 1418, 1420, 1425, 1427, 1429, 1431, 1433, 1435, 1437, 1439, 1442, 1444, 1446, 1448, 1450, 1452, 1458, 1461, 1464, 1467, 1472, 1475, 1478, 1481, 1487, 1490, 1493, 1496, 1500, 1503, 1506, 1509, 1511, 1513, 1516, 1520, 1523, 1528, 1531, 1534, 1536, 1539, 1542, 1545, 1548, 1551, 1555, 1558, 1561, 1564, 1569, 1572, 1575, 1577, 1580, 1584, 1587, 1589, 1591, 1594, 1599, 1602, 1605, 1608, 1611, 1614, 1618, 1621, 1624, 1627, 1630, 1633, 1636, 1639, 1642, 1645, 1648, 1652, 1655, 1658, 1661, 1663, 1665, 1668, 1671, 1674, 1676, 1681, 1685, 1687, 1690, 1693, 1695, 1698, 1700, 1702, 1704, 1707, 1710, 1714, 1717, 1719, 1722, 1725, 1727, 1730, 1733, 1735, 1738, 1741, 1745, 1748, 1750, 1753, 1755, 1758, 1761, 1764, 1766, 1769, 1772, 1774, 1778, 1781, 1784, 1787, 1790, 1793, 1798, 1801, 1804, 1807, 1811, 1814, 1817, 1819, 1822, 1825, 1828, 1831, 1834, 1837, 1840, 1844, 1847, 1850, 1853, 1856, 1861, 1864, 1867, 1870, 1873, 1877, 1882, 1885, 1888, 1891, 1894, 1897, 1900, 1903, 1911, 1915, 1919, 1923, 1927, 1931, 1936, 1940, 1946, 1950, 1954, 1958, 1963, 1967, 1971, 1975, 1979, 1985, 1990, 1994, 2000, 2004, 2010, 2017, 2021, 2027, 2038, 2040, 2043, 2048, 2051, 2054, 2059, 2065, 2067, 2069, 2071, 2073, 2075, 2091, 2097, 2112, 2117, 2124, 2134, 2137, 2142, 2148, 2153, 2166, 2178, 2193, 2205, 2222, 2224, 2229, 2231, 2236, 2242, 2247, 2253, 2259, 2265, 2272, 2280, 2289, 2305, 2310, 2315, 2320, 2326, 2331, 2336, 2341, 2346, 2353, 2360, 2367, 2372, 2378, 2383, 2388, 2393, 2400, 2410, 2418, 2428, 2437, 2445, 2455, 2463, 2473, 2483, 2490, 2499, 2504, 2510, 2518, 2526, 2534, 2541, 2550, 2557, 2564, 2571, 2578, 2586, 2595, 2601, 2607, 2613, 2619, 2625, 2632, 2638, 2644, 2651, 2656, 2662, 2669, 2675, 2682, 2689, 2699, 2706, 2717, 2725, 2736, 2744, 2754, 2761, 2770, 2779, 2786, 2796, 2805, 2811, 2821, 2828, 2833, 2837, 2841, 2847, 2853, 2858, 2866, 2873, 2880, 2886, 2893, 2901, 2909, 2915, 2921, 2929, 2937, 2945, 2953, 2962, 2971, 2984, 2993, 3002, 3016, 3024, 3033, 3051, 3063, 3071, 3079, 3090, 3098, 3106, 3114, 3124, 3133, 3141, 3153, 3164, 3173, 3178, 3187, 3200, 3216, 3229, 3235, 3247, 3257, 3262, 3268, 3287, 3296, 3305, 3316, 3324, 3332, 3340, 3349, 3357, 3365, 3372, 3381, 3389, 3398, 3406, 3415, 3425, 3433, 3443, 3453, 3462, 3472, 3481, 3497, 3502, 3507, 3514, 3521, 3529, 3536, 3541, 3549, 3554, 3561, 3569, 3575, 3578, 3587, 3596, 3601, 3606, 3613, 3622, 3628, 3633, 3640, 3646, 3653, 3660, 3667, 3672, 3680, 3683, 3690, 3695, 3700, 3705, 3710, 3715, 3721, 3727, 3735, 3744, 3750, 3756, 3763, 3769, 3776, 3784, 3790, 3795, 3800, 3805, 3810, 3816, 3821, 3828, 3837, 3847, 3856, 3865, 3871, 3881, 3887, 3896, 3904, 3913, 3917, 3920, 3925, 3931, 3944, 3946, 3948, 3950, 3952, 3954, 3956, 3959, 3961, 3963, 3966, 3969, 3971, 3973, 3975, 3977, 3982, 3987, 3989, 3992, 3994, 4000, 4002, 4004, 4007, 4010, 4012, 4015, 4018, 4021, 4024, 4029, 4034, 4036, 4038, 4040, 4042, 4044, 4047, 4049, 4051, 4053, 4056, 4059, 4064, 4069, 4074, 4079, 4082, 4088, 4093, 4096, 4099, 4102, 4105, 4112, 4115, 4118, 4120, 4123, 4126, 4129, 4132, 4135, 4138, 4144, 4147, 4150, 4152, 4154, 4157, 4162, 4164, 4168, 4171, 4175, 4178, 4181, 4184, 4187, 4190, 4193, 4196, 4199, 4202, 4206, 4209, 4212, 4215, 4218, 4223, 4225, 4227, 4230, 4233, 4239, 4241, 4243, 4245, 4247, 4250, 4253, 4255, 4258, 4261, 4265, 4268, 4271, 4274, 4277, 4280, 4283, 4286, 4291, 4294, 4298, 4301, 4304, 4307, 4310, 4313, 4316, 4319, 4322, 4325, 4329, 4332, 4335, 4338, 4341, 4348, 4354, 4357, 4363, 4367, 4372, 4378, 4382, 4386, 4392, 4399, 4405, 4411, 4418, 4423, 4428, 4432, 4436, 4446, 4452, 4457, 4462, 4468, 4471, 4474, 4477, 4479, 4481, 4484, 4487, 4493, 4498, 4503, 4508, 4513, 4516, 4521, 4524, 4527, 4530, 4532, 4544, 4558, 4573, 4578, 4583, 4588, 4594, 4602, 4607, 4613, 4620, 4626, 4631, 4636, 4641, 4646, 4652, 4657, 4662, 4667, 4673, 4678, 4685, 4695, 4704, 4712, 4718, 4723, 4730, 4739, 4745, 4753, 4758, 4766, 4773, 4779, 4784, 4791, 4799, 4808, 4816, 4821, 4826, 4831, 4842, 4848, 4858, 4865, 4871, 4877, 4883, 4889, 4896, 4902, 4909, 4913, 4920, 4928, 4934, 4941, 4951, 4961, 4971, 4981, 4991, 5001, 5010, 5021, 5029, 5039, 5051, 5061, 5071, 5080, 5088, 5095, 5103, 5114, 5120, 5125, 5129, 5134, 5138, 5144, 5151, 5158, 5166, 5173, 5180, 5186, 5191, 5203, 5214, 5219, 5224, 5232, 5239, 5250, 5263, 5269, 5278, 5288, 5294, 5301, 5308, 5321, 5330, 5340, 5349, 5357, 5368, 5376, 5383, 5390, 5400, 5407, 5414, 5422, 5431, 5441, 5452, 5462, 5471, 5482, 5498, 5506, 5521, 5528, 5543, 5554, 5563, 5573, 5582, 5591, 5600, 5606, 5614, 5624, 5629, 5639, 5644, 5655, 5670, 5675, 5682, 5687, 5694, 5705, 5715, 5724, 5740, 5751, 5765, 5780, 5789, 5799, 5807, 5816, 5826, 5829, 5836, 5840, 5849, 5854, 5863, 5880, 5889, 5893, 5899, 5916, 5928, 5941, 5945, 5951, 5960, 5969, 5978, 5985, 5995, 6004, 6010, 6018, 6030, 6034, 6042, 6049, 6057, 6061, 6069, 6081, 6087, 6099, 6104, 6116, 6127, 6138, 6150, 6162, 6174, 6186, 6197, 6202, 6213, 6226, 6234, 6243, 6252, 6263, 6274, 6284, 6293, 6303, 6315, 6324, 6331, 6339, 6346, 6352, 6360, 6367, 6374, 6379, 6385, 6394, 6402, 6410, 6417, 6424, 6431, 6438, 6442, 6451, 6462, 6471, 6481, 6490, 6499, 6507, 6517, 6523, 6532, 6540, 6546, 6554, 6563, 6571, 6580, 6586, 6592, 6601, 6612, 6623, 6631, 6639, 6644, 6651, 6657, 6664, 6671, 6680, 6688, 6697, 6703, 6714, 6723, 6735, 6747, 6757, 6764, 6772, 6778, 6783, 6788, 6800, 6803, 6806, 6809, 6812, 6815, 6818, 6820, 6826, 6829, 6832, 6834, 6837, 6841, 6843, 6845, 6847, 6850, 6853, 6857, 6860, 6863, 6866, 6869, 6872, 6875, 6877, 6880, 6883, 6887, 6890, 6893, 6896, 6901, 6904, 6907, 6910, 6913, 6917, 6920, 6923, 6926, 6929, 6932, 6935, 6938, 6941, 6945, 6948, 6951, 6954, 6957, 6960, 6963, 6966, 6969, 6972, 6976, 6979, 6982, 6984, 6987, 6989, 6992, 6994, 6996, 6998, 7000, 7003, 7006, 7009, 7012, 7015, 7018, 7021, 7024, 7027, 7030, 7034, 7037, 7040, 7043, 7046, 7049, 7054, 7057, 7060, 7064, 7067, 7070, 7073, 7076, 7080, 7083, 7086, 7089, 7092, 7098, 7101, 7104, 7107, 7110, 7115, 7120, 7123, 7126, 7130, 7133, 7138, 7141, 7144, 7147, 7149, 7151, 7153, 7155, 7158, 7160, 7162, 7164, 7171, 7174, 7176, 7178, 7181, 7183, 7185, 7187, 7189, 7192, 7194, 7198, 7201, 7204, 7207, 7209, 7212, 7215, 7220, 7223, 7228, 7231, 7233, 7238, 7243, 7248, 7253, 7259, 7264, 7269, 7274, 7278, 7284, 7290, 7295, 7297, 7299, 7301, 7304, 7307, 7311, 7314, 7317, 7320, 7325, 7328, 7331, 7334, 7336, 7338, 7341, 7347, 7349, 7351, 7354, 7357, 7360, 7362, 7365, 7367, 7369, 7372, 7376, 7379, 7382, 7385, 7387, 7390, 7393, 7396, 7401, 7404, 7407, 7411, 7416, 7421, 7426, 7431, 7436, 7440, 7443, 7446, 7449, 7457, 7461, 7467, 7473, 7480, 7486, 7492, 7498, 7505, 7512, 7518, 7523, 7530, 7536, 7542, 7553, 7558, 7566, 7574, 7578, 7581, 7584, 7587, 7592, 7595, 7600, 7603, 7609, 7614, 7619, 7624, 7629, 7632, 7635, 7651, 7656, 7661, 7666, 7670, 7675, 7680, 7685, 7690, 7695, 7702, 7709, 7714, 7717, 7722, 7727, 7730, 7736, 7750, 7764, 7782, 7797, 7804, 7815, 7820, 7823, 7828, 7838, 7847, 7858, 7863, 7868, 7872, 7881, 7890, 7896, 7906, 7911, 7916, 7922, 7932, 7941, 7950, 7959, 7964, 7969, 7976, 7983, 7989, 7994, 7999, 8008, 8013, 8018, 8023, 8029, 8034, 8040, 8045, 8050, 8056, 8064, 8069, 8076, 8081, 8087, 8092, 8097, 8104, 8113, 8119, 8128, 8137, 8142, 8153, 8163, 8172, 8182, 8192, 8198, 8204, 8215, 8223, 8231, 8240, 8248, 8254, 8261, 8270, 8283, 8295, 8301, 8312, 8324, 8335, 8342, 8349, 8358, 8371, 8381, 8391, 8398, 8412, 8426, 8431, 8437, 8450, 8457, 8467, 8478, 8484, 8493, 8501, 8506, 8514, 8523, 8531, 8538, 8545, 8554, 8562, 8569, 8577, 8583, 8590, 8598, 8608, 8624, 8635, 8648, 8651, 8658, 8663, 8669, 8679, 8689, 8695, 8702, 8707, 8716, 8728, 8740, 8750, 8755, 8762, 8767, 8778, 8797, 8815, 8824, 8828, 8838, 8848, 8861, 8873, 8884, 8894, 8912, 8930, 8955, 8964, 8976, 8981, 8990, 9000, 9012, 9017, 9026, 9037, 9045, 9054, 9063, 9072, 9083, 9095, 9100, 9109, 9119, 9125, 9133, 9145, 9150, 9155, 9165, 9175, 9184, 9194, 9199, 9210, 9215, 9220, 9230, 9238, 9248, 9253, 9260, 9265, 9274, 9284, 9290, 9299, 9305, 9311, 9319, 9327, 9332, 9338, 9345, 9354, 9360, 9365, 9371, 9376, 9381, 9386, 9391, 9396, 9401, 9408, 9414, 9419, 9425, 9430, 9441, 9446, 9453, 9464, 9473]
all_questions = get_indexes_of_questions()
print(len(all_questions))
all_answers = ['3', '1', '4', '2', '2', '2', '3', '1', '3', '4', '3', '4', '4', '4', '2', '3', '2', '2', '3', '3', '2', '3', '4', '3', '2', '1', '2', '3', '2', '2', '2', '2', '4', '1', '1', '2', '4', '2', '2', '4', '4', '3', '1', '4', '4', '3', '1', '2', '4', '2', '1, 2, 5', '3, 4, 7', '1, 4, 7', '3, 5, 7', '1, 4, 5', '1, 3, 5', '4', '3', '2', '3, 6', '1, 6', '4', '3', '2, 5', '1, 6, 7', '1, 6, 7', '1, 3, 6', '1, 3, 6', '2, 3, 6', '1, 4, 5', '1, 4, 5', '2', '1, 2, 3', '4, 6, 7', '1, 5, 7', '5, 6, 7', '3', '4', '1', '2', '3', '4', '1', '2', '3', '4', '1', '3', '4', '2', '1', '3', '3', '2', '4', '1', '3', '3', '1', '3', '2', '4', '2', '1', '2', '5', '2', '4', '3', '5, 7', '5, 6', '4, 6', '1', '2', '4', '1', '651324', '     526341', '253614', '521364', '613524', '235146', '     325146', '321654', '634521', '651324', '264351', '     352416', '235641', '634521', '241635', '534621', '     632451', '316542', '156324', '512463', '524163', '     523614', '346251', '316425', '345126', '651432', '     523461', '2', '2', '4', '4', '1', '4', '1', '1', '2', '4', '2', '1', '4', '4', '2', '4', '4', '4', '4', '735621', '     674512', '736124', '457213', '537241', '421653', '     654172', '356714', '  735621', '761245', '751346', '     642715', '612753', '257361', '562731', '546231', '     421376', '674512', '437251', '437251', '365712', '     275361', '261347', '531726', '476231', '174523', '     614752', '473512', '2', '3', '4', '2', '2, 3, 4', '1', '2, 4', '3, 7, 8', '2, 5', '1, 5, 8', '3, 5', '2, 5, 7', '2', '1, 5, 7', '2, 3, 5', '4', '3', '2, 4', '2, 6, 7', '2', '2, 4, 6', '3', '4', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '     0', '1', '4', '2', '2', '1', '3', '1', '4', '1', '3', '2', '4', '2', '2', '1', '2', '1', '4', '2', '2', '2', '3', '2', '1', '3', '1', '1', '4', '1', '2', '1', '4', '2', '1', '4', '2', '3', '2', '3', '3', '1', '3', '2', '2', '1', '3', '2', '3', '2', '3', '1', '2', '4', '4', '2', '2', '2', '2', '2', '2', '2', '4', '2', '2', '3', '2', '3', '3', '4', '3', '4', '4', '1', '4', '3', '3', '3', '3', '3', '4', '1', '3', '4', '1', '1', '2', '2, 5, 7', '1, 4, 6', '3', '1, 2, 4', '2, 5, 8', '1, 3, 4', '1, 3, 6', '1, 2, 6', '1, 5, 8', '1, 3, 5', '4', '3', '4', '3, 4, 5', '2', '3', '2', '2', '1', '1', '1, 4, 5', '4', '2, 3, 6', '4', '2', '1', '1', '1', '2', '1, 4, 5', '3', '2', '2', '1, 3, 7', '3', '1, 4, 7', '2, 3, 5', '3, 6, 8', '4', '2', '3', '1', '3', '1', '3', '3', '1', '2', '1', '3, 6', '4, 5', '4, 7', '564231', '     452136', '612453', '532164', '526413', '362514', '     356214', '351642', '     234156', '631542', '132456', '643152', '213654', '     153462', '261435', '     263514', '642531', '563142', '542136', '463521', '     532641', '531264', '     652143', '264531', '245631', '315426', '652143', '     653214', '256314', '     264351', '534162', '243615', '1', '1', '1', '4', '2', '4', '4', '4', '1', '4', '4', '2', '2', '4', '2', '4', '4', '4', '4', '451237', '     736124', '421376', '657231', '572431', '417253', '     257134', '134567', '     764315', '67123', '567213', '542371', '457123', '     246173', '631725', '     251673', '456731', '421376', '375124', '351742', '     764325', '354672', '     246175', '741253', '364251', '647253', '651374', '1', '1, 4', '2, 5', '4, 5, 7', '3', '3', '2,4', '3,4', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '4', '2', '4', '1', '3', '3', '1', '3', '1', '4', '2', '2', '1', '2', '1', '2', '1', '2', '4', '4', '2', '1', '2', '3', '4', '2', '3', '3', '3', '3', '2', '2', '3', '2', '2', '2', '1', '3', '2', '4', '1', '2', '4', '2', '4', '2', '3', '2', '2', '2', '1', '3', '3', '4', '4', '4', '3', '2', '4', '4', '2', '2', '1', '3', '1', '2', '4', '1', '1', '3', '1', '1', '4', '1', '1', '2', '2', '3', '2', '1', '4', '1', '2', '1', '3', '4', '3', '2', '1', '3', '4', '2', '2', '4', '1, 3, 6', '2, 5, 7', '2, 3, 6', '4, 5, 7', '1, 3, 7', '1, 4, 5', '1, 3, 5', '4, 5', '4', '1', '4', '4', '2', '1, 2', '2', '4', '3', '3', '1, 4', '1, 3', '3', '2', '4', '1', '2', '2', '3, 5, 7', '1, 2, 4', '2, 4, 6', '2', '1, 4, 5', '2, 3, 7', '5, 6, 7', '2, 3, 5', '3, 4, 6', '4', '2', '3', '3, 5, 7', '4', '2', '2', '2, 3', '3', '2', '4', '2', '4, 7', '4, 8', '5, 6', '3', '1', '4', '2', '1', '314526', '     361254', '325641', '251436', '421365', '541236', '     326145', '362514', '     356214', '463215', '124563', '254316', '436215', '     241536', '431256', '     214365', '142536', '153264', '361452', '435621', '     536412', '165432', '     621543', '651432', '145632', '326451', '3', '2', '4', '2', '4', '1', '2', '3', '1', '4', '1', '4', '3', '1', '2', '2', '4', '4', '4', '2', '1', '4', '4', '4', '4', '2', '4', '1', '2', '2', '347516', '     742531', '462753', '276143', '361275', '243617', '     361524', '162753', '     423517', '741625', '517362', '451632', '675324', '     432157', '417652', '     135742', '572163', '514726', '376251', '147236', '     516472', '341257', '     672134', '241736', '561372', '751634', '647123', '     256731', '142657', '     654132', '534176', '512437', '653142', '3, 5', '2, 3, 5', '1', '4', '2, 3, 5', '4', '3, 5, 7', '1, 2, 3', '2, 4', '2, 5, 6', '4, 5', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '3', '3', '4', '1', '3', '3', '4', '2', '1', '2', '4', '1', '2', '3', '1', '2', '2', '3', '1', '4', '2', '1', '2', '3', '3', '2', '4', '3', '1', '4', '2', '4', '3', '2', '4', '3', '4', '1', '2', '2', '3', '3', '3', '3', '3', '2', '4', '1', '3', '2', '2', '2', '3', '1', '4', '3', '3', '4', '2', '3', '4', '1', '4', '1', '3', '1, 5, 7', '2, 4, 8', '2, 5, 7', '1, 3, 6', '2, 3, 5', '1, 4, 5', '1, 4, 7', '1, 3, 6', '2, 5, 6', '1, 3, 5', '2, 3, 7', '2, 5, 6', '2', '1', '4', '1, 2, 7', '2, 4, 7', '2, 4, 8', '2, 4, 7', '3, 5, 6', '1', '1', '4', '2', '3', '4', '2, 3, 5', '1, 2, 4', '3', '1, 2, 4', '2', '3', '4', '1, 5, 7', '3, 5, 7', '4', '1, 2, 3', '1, 4, 6', '1, 4, 7', '4, 5, 7', '2, 3, 5', '2', '4', '2', '3', '4', '4', '3', '2', '4', '2', '4', '1', '2', '2', '3', '4', '1', '2', '2', '1', '3', '4', '2', '4', '3', '4', '6', '3,7', '3, 7', '3, 7', '4, 5', '3, 5', '5, 9', '3, 4', '1', '4', '5', '463125', '    526431', '352641', '352641', '635142', '264153', '     651423', '243615', '     163425', '125634', '614235', '462135', '164352', '     351246', '563214', '     245361', '635214', '415362', '432651', '461253', '     423561', '461325', '     162435', '325461', '356421', '645213', '163524', '136524', '243561', '    261435', '513264', '326415', '1', '4', '2', '4', '1', '4', '4', '461253', '     612753', '612357', '671423', '437162', '745623', '     246173', '457132', '     237154', '354172', '546731', '631257', '361725', '     174256', '614527', '     257641', '743652', '246173', '257641', '274136', '     364712', '361275', '     314276', '641325', '531724', '2', '2', '2, 4', '4, 6', '2', '2', '5, 8, 9', '2, 5, 7', '3, 5, 6', '3, 4, 6', '4, 5, 6', '3, 5, 6', '2, 6, 7', '2, 5, 6', '5, 7, 8', '2, 4, 6', '3, 5, 6', '3, 4, 6', '3, 5, 7', '3, 4, 6', '2, 3, 4', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0', '0', '     0', '0', '0', '0', '0', '     0']

print ("length of questions ")
print (len(all_questions))
print ("answers are {}".format(len(all_answers)))
if len(all_answers) != len(all_questions):
    print ("FUCCCCK")
else:
    print ("perhaps URAAAA")
# print (all_questions)
def add_answers(all_questions, all_answers):
    """
    iterates throgh file and adds answers in blue next to question
    """
    for i in range(len(all_questions)):
        if i%100 == 0:
            print (i)
        actual_question = patm.paragraphs[all_questions[i]]
        # # code below should change the color
        run = actual_question.add_run("  "  + str(all_answers[i]))
        # actual_question.text += ("  "  + str(dummy_answers[i]))
        font = run.font
        font.color.rgb = RGBColor(0x00, 0x00, 0xff)
        font.size = Pt(15)

    patm.save("patm_3_without_chisht_sxal.docx")
    print ("Done adding answers")

add_answers(all_questions, all_answers)
#//////////////////////////////////////////////////////////////////////////////

# import numpy as np
# import docx

# patm_tables = docx.Document(
#     r"C:\Users\Hayk\Desktop\patm qnnutun patasxannerov\3-րդ շտեմարան\Պատասխաններ.docx")
# tables = patm_tables.tables


# def get_answers(table):
#     """
#     Function takes as input table and return a 1d array of answers
#     """
#     answers = []
#     row_count = len(table.rows)
#     col_count = len(table.rows[0].cells)

#     for row in table.rows[2:]:  # 2: becaouse of text on the top
#         for counter, cell in zip(range(col_count), row.cells):
#             if counter % 2 == 1:
#                 print()
#                 for paragraph in cell.paragraphs:
#                     answers.append(paragraph.text)

#     answers = np.array(answers)
#     answers = answers.reshape(row_count - 2, col_count // 2)
#     answers = answers.T
#     answers = answers.reshape(1, (row_count - 2) * col_count // 2)

#     answers = list(filter(lambda x: x not in [
#                    "", " ", "  ", "   ", "    ", "     ", "      "], answers[0]))

#     return answers


# all_answers = []
# for i in range(len(tables)):

#     all_answers += get_answers(tables[i])

# print(len(all_answers))
# print(all_answers)
# print("first 2")
# print(all_answers[0])
# print(all_answers[1])
# print("last_two")
# print(all_answers[-1])
# print(all_answers[-2])
# PENTAGON .DESTROY
# answers = answers[24:]